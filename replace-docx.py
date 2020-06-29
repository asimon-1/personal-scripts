"""
To use, enter a command prompt in the same directory where this file lives and
run the program. You will be prompted for the file path, a regex pattern of text
to replace, and the replacement text to enter.
You can use the --inplace cmd flag to overwrite the existing file. Only do this
if you are absolutely sure!
Using code provided in https://stackoverflow.com/a/42829667
Requires the python-docx module.
"""

import re
import sys

from docx import Document


def docx_replace_regex(doc_obj, regex, replace):
    """Find regex pattern within document and replace it

    Arguments:
        doc_obj {docx.Document} -- The document object to be manipulated
        regex {str} -- Regular expression pattern of text to be replaced
        replace {str} -- New text to replace the regex matches
    """
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex, replace)


if __name__ == "__main__":
    try:
        inplace = sys.argv[1] == "--inplace"
    except:
        inplace = False

    filename = input("Please enter the path to the file including filename: ")
    pattern = input("Please enter the regular expression of text to replace: ")
    replacement = input("Please enter the replacement text: ")
    if inplace:
        new_filename = filename
    else:
        new_filename = filename[:-5] + "_replaced.docx"

    regex = re.compile(pattern)
    doc = Document(filename)
    docx_replace_regex(doc, regex, replacement)
    doc.save(new_filename)
