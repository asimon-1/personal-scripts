"""Imports requirements from an SRS document and puts them in an Excel sheet for easy manipulation.
> python import-from-srs.py source-srs.docx destination.xlsx
"""

import re
import sys
import xlsxwriter

from docx import Document


def docx_find_regex(doc_obj, regex):
    """Find regex pattern within document and return all paragraphs with matches.

    Arguments:
        doc_obj {docx.Document} -- The document object
        regex {str} -- Regular expression pattern of text to be found
    """
    hits = [p.text.strip() for p in doc_obj.paragraphs if regex.search(p.text)]

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                for cell_hit in docx_find_regex(cell, regex):
                    hits.append(cell_hit)

    return hits


if __name__ == "__main__":
    try:
        source = sys.argv[1]
        destination = sys.argv[2]
    except:
        raise Exception("Incorrect number of arguments.")

    pattern = r"R \[\w{4}\]"
    regex = re.compile(pattern)
    doc = Document(source)
    hits = docx_find_regex(doc, regex)

    workbook = xlsxwriter.Workbook(destination)
    worksheet = workbook.add_worksheet()
    header_rows = 1

    worksheet.write(0, 0, "Raw Text")
    worksheet.write(0, 1, "Requirement Number")
    worksheet.write(0, 2, "Requirement Text")
    for ind, hit in enumerate(hits):
        req_num = hit.split("]", 1)[0].split("[", 1)[1]
        req_text = hit.split("]", 1)[1]
        worksheet.write(header_rows + ind, 0, hit)
        worksheet.write(header_rows + ind, 1, req_num)
        worksheet.write(header_rows + ind, 2, req_text)
    workbook.close()
